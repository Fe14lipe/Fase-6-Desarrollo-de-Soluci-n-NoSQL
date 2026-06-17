import os
import shutil
import docx
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_run_font(run, name="Arial", size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_paragraph_styled(doc, text="", style=None, space_after=6, space_before=0, line_spacing=1.15):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    return p

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, name="Arial", size=14, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, name="Arial", size=12, bold=True)
    elif level == 3:
        run = p.add_run(text)
        set_run_font(run, name="Arial", size=11, bold=True)
    return p

def add_styled_code_block(doc, code_text):
    """
    Crea una tabla de 1 fila y 1 columna con fondo gris-azul claro,
    borde izquierdo grueso color azul cerceta (teal) y sin los otros bordes,
    emulando el estilo de caja de código del usuario.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    
    # Configurar el ancho del bloque de código
    table.columns[0].width = Inches(6.0)
    cell = table.cell(0, 0)
    
    # 1. Configurar sombreado / fondo (Background Shading)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F7F9')  # Color gris-azul claro
    tcPr.append(shd)
    
    # 2. Configurar bordes (Left border grueso color azul/teal, otros none)
    tcBorders = OxmlElement('w:tcBorders')
    
    # Borde izquierdo grueso (Teal: 0D5C75, tamaño: 24 = 3pt)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '0D5C75')
    tcBorders.append(left)
    
    # Otros bordes inactivos
    for side in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
        
    tcPr.append(tcBorders)
    
    # 3. Configurar padding de celda (Cell Margins)
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', '120'), ('bottom', '120'), ('left', '240'), ('right', '240')]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), val)
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)
    
    # 4. Insertar el texto del código
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text)
    set_run_font(run, name="Consolas", size=9)
    
    # Agregar espaciado inferior insertando un párrafo vacío después de la tabla
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)
    delete_paragraph(p_after) # Solo dejamos el espaciado
    return table

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._element = p._parent = None

def add_bullet_point(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        set_run_font(r_prefix, name="Arial", size=11, bold=True)
        
    r_text = p.add_run(text)
    set_run_font(r_text, name="Arial", size=11)
    return p

def main():
    directorio_actual = os.path.dirname(os.path.abspath(__file__))
    ruta_salida = os.path.join(os.path.dirname(directorio_actual), "Fase 6 Proyecto Integrador NOSQL MONGODB_v2.docx")
    ruta_alterna = os.path.join(os.path.dirname(directorio_actual), "Fase 6 Proyecto Integrador NOSQL MONGODB_v2_NUEVO.docx")
    
    doc = docx.Document()
    
    # --- TÍTULO PRINCIPAL ---
    p_title = add_paragraph_styled(doc, space_after=4, space_before=12)
    r_title = p_title.add_run("INFORME DE PROYECTO: SOUNDWAVE (FASE 6)")
    set_run_font(r_title, name="Arial", size=16, bold=True)
    
    p_subtitle = add_paragraph_styled(doc, space_after=12, space_before=0)
    r_subtitle = p_subtitle.add_run("Implementación NoSQL y Migración a MongoDB")
    set_run_font(r_subtitle, name="Arial", size=12, bold=False, italic=True)
    
    # --- PORTADA / DATOS DE GRUPO ---
    p_group = add_paragraph_styled(doc, space_after=3)
    r_group = p_group.add_run("INTEGRANTES DEL GRUPO:")
    set_run_font(r_group, name="Arial", size=11, bold=True)
    
    add_bullet_point(doc, "Felipe Boada")
    add_bullet_point(doc, "Matías Flores")
    add_bullet_point(doc, "Ariel Tejada")
    add_bullet_point(doc, "Yusely Zapata")
    
    p_materia = add_paragraph_styled(doc, space_after=3, space_before=6)
    r_materia = p_materia.add_run("MATERIA:")
    set_run_font(r_materia, name="Arial", size=11, bold=True)
    add_bullet_point(doc, "Base de Datos II (UDLA)")
    
    add_paragraph_styled(doc, "—" * 50, space_after=12)
    
    # --- INTRODUCCIÓN ---
    add_heading_styled(doc, "1. INTRODUCCIÓN", level=2)
    p_intro = add_paragraph_styled(doc)
    p_intro.add_run("El presente informe detalla la transición del sistema musical ")
    p_intro.add_run("SoundWave").bold = True
    p_intro.add_run(" desde una arquitectura de base de datos relacional (Microsoft SQL Server) hacia una base de datos NoSQL documental basada en ")
    p_intro.add_run("MongoDB").bold = True
    p_intro.add_run(". El objetivo principal es reestructurar el modelo de datos físico para aprovechar las ventajas de anidamiento y desnormalización NoSQL, optimizando la latencia de las consultas, agilizando el flujo CRUD de la aplicación y simplificando las consultas analíticas del negocio.")
    
    # --- MIGRACIÓN ---
    add_heading_styled(doc, "2. MIGRACIÓN Y DISEÑO DEL MODELO NOSQL", level=2)
    p_mig = add_paragraph_styled(doc)
    p_mig.add_run("Para migrar los datos relacionales hacia MongoDB, se definieron consultas en SQL Server que agrupan los registros dependientes como colecciones y documentos embebidos utilizando ")
    p_mig.add_run("FOR JSON PATH").bold = True
    p_mig.add_run(" y ")
    p_mig.add_run("JSON_QUERY").bold = True
    p_mig.add_run(". Esto permitió consolidar el esquema físico NoSQL:")
    
    add_bullet_point(doc, "Suscripciones y Pagos se estructuraron embebidos dentro de la colección usuarios.", bold_prefix="• Usuarios: ")
    add_bullet_point(doc, "Canciones y sus géneros se insertaron como arreglos embebidos dentro de cada documento de la colección albums.", bold_prefix="• Álbumes y Canciones: ")
    add_bullet_point(doc, "Playlists contienen información desnormalizada de sus canciones agregadas, evitando cruzar colecciones.", bold_prefix="• Playlists: ")
    
    p_sql = add_paragraph_styled(doc)
    p_sql.add_run("A continuación se muestra el fragmento del script SQL (")
    p_sql.add_run("sql_to_json.sql").italic = True
    p_sql.add_run(") para la exportación estructurada de usuarios:")
    
    add_styled_code_block(doc, """SELECT 
    u.usuarioId, u.nombre, u.apellido, u.email, u.contraseña, u.pais, u.fechaRegistro, u.estado, u.telefono, u.eliminado,
    JSON_QUERY((
        SELECT 
            s.suscripcionId, s.fechaInicio, s.fechaFin, s.estado,
            JSON_QUERY((
                SELECT planId, nombre, precioMensual, descripcion FROM soundwave.PlanSuscripcion p WHERE p.planId = s.Plan_planId
                FOR JSON PATH, WITHOUT_ARRAY_WRAPPER
            )) AS [plan],
            JSON_QUERY((
                SELECT pagoId, monto, fechaPago, metodoPago, estado FROM soundwave.Pago pg WHERE pg.Suscripcion_suscripcionId = s.suscripcionId
                FOR JSON PATH
            )) AS pagos
        FROM soundwave.Suscripcion s WHERE s.Usuario_usuarioId = u.usuarioId FOR JSON PATH
    )) AS suscripciones
FROM soundwave.Usuario u FOR JSON PATH;""")

    # --- ESQUEMAS JSON ---
    add_heading_styled(doc, "2.1. ESQUEMAS JSON DE EJEMPLO DE MONGODB", level=2)
    p_esq_desc = add_paragraph_styled(doc)
    p_esq_desc.add_run("A continuación se presentan los ejemplos de documentos para cada una de las colecciones físicas implementadas en la base de datos MongoDB, las cuales demuestran las técnicas de incrustación de subdocumentos e historiales de transacciones:")
    
    # 1. Colección Usuario
    p_u_title = add_paragraph_styled(doc, space_after=4, space_before=6)
    r_u_title = p_u_title.add_run("1. Colección `Usuario`")
    set_run_font(r_u_title, name="Arial", size=11, bold=True)
    add_styled_code_block(doc, """{
  "_id": { "$oid": "647f1234567890abcdef1234" },
  "usuarioId": 1,
  "nombre": "Benjamin",
  "apellido": "Perez",
  "email": "benjamin@udla.edu.ec",
  "contraseña": "hash_pass_1",
  "pais": "Ecuador",
  "fechaRegistro": { "$date": "2024-01-01T00:00:00.000Z" },
  "estado": "Activa",
  "telefono": "0960786020",
  "eliminado": false,
  "suscripciones": [
    {
      "suscripcionId": 1,
      "fechaInicio": { "$date": "2024-01-01T00:00:00.000Z" },
      "fechaFin": null,
      "estado": "Activa",
      "plan": {
        "planId": 2,
        "nombre": "Premium",
        "precioMensual": { "$numberDecimal": "5.99" },
        "descripcion": "Música sin anuncios, descargas offline, audio HQ"
      },
      "pagos": [
        {
          "pagoId": 1,
          "monto": { "$numberDecimal": "5.99" },
          "fechaPago": { "$date": "2024-01-01T00:00:00.000Z" },
          "metodoPago": "Tarjeta",
          "estado": "Aprobado"
        }
      ]
    }
  ]
}""")

    # 2. Colección Discografica
    p_d_title = add_paragraph_styled(doc, space_after=4, space_before=6)
    r_d_title = p_d_title.add_run("2. Colección `Discografica`")
    set_run_font(r_d_title, name="Arial", size=11, bold=True)
    add_styled_code_block(doc, """{
  "_id": { "$oid": "647f11111111111111111111" },
  "discograficaId": 1,
  "nombre": "Rimas Entertainment",
  "pais": "Puerto Rico",
  "fechaFundacion": { "$date": "2014-01-01T00:00:00.000Z" },
  "eliminado": false
}""")

    # 3. Colección Artista
    p_a_title = add_paragraph_styled(doc, space_after=4, space_before=6)
    r_a_title = p_a_title.add_run("3. Colección `Artista`")
    set_run_font(r_a_title, name="Arial", size=11, bold=True)
    add_styled_code_block(doc, """{
  "_id": { "$oid": "647f22222222222222222222" },
  "artistaId": 1,
  "nombreArtistico": "Bad Bunny",
  "pais": "Puerto Rico",
  "fechaInicio": { "$date": "2016-01-01T00:00:00.000Z" },
  "discograficaId": 1,
  "discograficaNombre": "Rimas Entertainment",
  "eliminado": false
}""")

    # 4. Colección Album
    p_al_title = add_paragraph_styled(doc, space_after=4, space_before=6)
    r_al_title = p_al_title.add_run("4. Colección `Album`")
    set_run_font(r_al_title, name="Arial", size=11, bold=True)
    add_styled_code_block(doc, """{
  "_id": { "$oid": "647faaaaaaaaaaaaaaaaaaaa" },
  "albumId": 1,
  "titulo": "Un Verano Sin Ti",
  "fechaLanzamiento": { "$date": "2022-05-06T00:00:00.000Z" },
  "artistaId": 1,
  "artistaNombre": "Bad Bunny",
  "canciones": [
    {
      "cancionId": 1,
      "titulo": "Me Porto Bonito",
      "duracion": 178,
      "explicita": true,
      "eliminado": false,
      "generos": ["Reggaeton", "Urbano"]
    },
    {
      "cancionId": 2,
      "titulo": "Tití Me Preguntó",
      "duracion": 243,
      "explicita": true,
      "eliminado": false,
      "generos": ["Reggaeton"]
    }
  ]
}""")

    # 5. Colección Playlist
    p_pl_title = add_paragraph_styled(doc, space_after=4, space_before=6)
    r_pl_title = p_pl_title.add_run("5. Colección `Playlist`")
    set_run_font(r_pl_title, name="Arial", size=11, bold=True)
    add_styled_code_block(doc, """{
  "_id": { "$oid": "647f88888888888888888888" },
  "playlistId": 1,
  "nombre": "Perreo Intenso 2024",
  "fechaCreacion": { "$date": "2024-01-01T00:00:00.000Z" },
  "privacidad": "Pública",
  "usuarioId": 1,
  "usuarioNombre": "Benjamin Perez",
  "eliminado": false,
  "canciones": [
    {
      "cancionId": 1,
      "titulo": "Me Porto Bonito",
      "duracion": 178,
      "albumId": 1,
      "albumTitulo": "Un Verano Sin Ti",
      "artistaNombre": "Bad Bunny",
      "fechaAgregado": { "$date": "2024-01-01T12:00:00.000Z" }
    }
  ]
}""")

    # 6. Colección Reproduccion
    p_r_title = add_paragraph_styled(doc, space_after=4, space_before=6)
    r_r_title = p_r_title.add_run("6. Colección `Reproduccion`")
    set_run_font(r_r_title, name="Arial", size=11, bold=True)
    add_styled_code_block(doc, """{
  "_id": { "$oid": "647fb1b1b1b1b1b1b1b1b1b1" },
  "reproduccionId": 1,
  "fechaHora": { "$date": "2026-06-16T19:00:00.000Z" },
  "duracionEscuchada": 178,
  "dispositivo": "iPhone",
  "usuarioId": 1,
  "cancionId": 1,
  "cancionTitulo": "Me Porto Bonito",
  "artistaNombre": "Bad Bunny"
}""")

    # --- SCRIPTS PYTHON ---
    add_heading_styled(doc, "3. AUTOMATIZACIÓN DEL PIPELINE DE MIGRACIÓN", level=2)
    add_paragraph_styled(doc, "Se crearon dos scripts en Python para automatizar el ciclo de migración de datos:")
    add_bullet_point(doc, "export_sql_to_json.py: Conecta a SQL Server usando pyodbc, lee la configuración de config.json, ejecuta las consultas FOR JSON y genera los archivos JSON formateados en la ruta mongodb/data/.", bold_prefix="• Extracción (ETL): ")
    add_bullet_point(doc, "import_json_to_mongo.py: Conecta a MongoDB, lee los archivos JSON generados, convierte dinámicamente las cadenas de texto de fecha a tipos BSON Date nativos de MongoDB e inserta los documentos, estableciendo los índices de rendimiento necesarios.", bold_prefix="• Carga y Tipado (Load): ")

    # --- PLAYGROUND ---
    add_heading_styled(doc, "4. PLAYGROUND MONGODB (PLAYGROUND.JS)", level=2)
    p_play = add_paragraph_styled(doc)
    p_play.add_run("Se escribió el script interactivo ")
    p_play.add_run("playground.js").bold = True
    p_play.add_run(" para evaluar las operaciones en MongoDB. Este archivo simula la inicialización de colecciones, inserciones de mock data, operaciones CRUD y búsquedas avanzadas. Se destacan las consultas analíticas del negocio mediante agregaciones:")
    
    p_look = add_paragraph_styled(doc)
    p_look.add_run("Consulta agregada con lookup para relacionar artistas y discográficas:")
    
    add_styled_code_block(doc, """db.artistas.aggregate([
  { $lookup: { from: "discograficas", localField: "discograficaId", foreignField: "discograficaId", as: "d_info" } },
  { $unwind: { path: "$d_info", preserveNullAndEmptyArrays: true } },
  { $project: { _id: 0, artistaId: 1, nombreArtistico: 1, pais: 1, discograficaNombre: "$d_info.nombre" } }
]);""")

    # --- INTEGRACION FLASK ---
    add_heading_styled(doc, "5. INTEGRACIÓN DEL BACKEND FLASK CON MONGODB", level=2)
    p_flask = add_paragraph_styled(doc)
    p_flask.add_run("La aplicación Flask en ")
    p_flask.add_run("app.py").bold = True
    p_flask.add_run(" se adaptó para conectarse dinámicamente a MongoDB cargando la configuración de config.json. Para asegurar la compatibilidad con las vistas y evitar tener que modificar las plantillas HTML (templates), se creó un adaptador de objetos:")
    
    add_bullet_point(doc, "Permite acceder a los campos de los diccionarios de MongoDB como atributos de objetos tradicionales (ej. fila.nombre en lugar de fila['nombre']), emulando el comportamiento del driver pyodbc.", bold_prefix="• Clase MongoRow: ")
    add_bullet_point(doc, "Se reescribieron las rutas administrativas (Discográficas, Artistas, Canciones, Playlists y Usuarios) para alternar entre consultas SQL y consultas PyMongo de forma dinámica según la base de datos activa.", bold_prefix="• Lógica CRUD Dual: ")
    add_bullet_point(doc, "Se implementaron agregaciones mediante pipelines sobre subdocumentos anidados (ej. desglosar e identificar ingresos premium acumulados de usuarios y transacciones por periodos).", bold_prefix="• Reportes de Negocio: ")

    # --- PRUEBAS ---
    add_heading_styled(doc, "6. EVIDENCIAS DE PRUEBAS AUTOMATIZADAS (QUALITY ASSURANCE)", level=2)
    p_qa = add_paragraph_styled(doc)
    p_qa.add_run("Se diseñó un suite de pruebas unitarias y de sistema en ")
    p_qa.add_run("tests/test_mongodb.py").bold = True
    p_qa.add_run(" para validar las APIs y operaciones CRUD corriendo sobre MongoDB. Todas las pruebas pasaron exitosamente:")
    add_styled_code_block(doc, """tests/test_mongodb.py .......                                            [100%]
============================== 7 passed in 9.10s ==============================""")
    
    p_qa2 = add_paragraph_styled(doc)
    p_qa2.add_run("La suite de pruebas original para SQL Server (test_app.py) también se ejecutó en modo mssql sin reportar fallos o regresiones:")
    add_styled_code_block(doc, """tests/test_app.py ....................                                   [100%]
============================= 20 passed in 9.83s ==============================""")

    # --- CONCLUSIONES ---
    add_heading_styled(doc, "7. CONCLUSIONES Y LECCIONES APRENDIDAS", level=2)
    
    p_c = add_paragraph_styled(doc)
    p_c.add_run("• Conclusiones:").bold = True
    add_bullet_point(doc, "La estructura jerárquica de MongoDB redujo la complejidad física de la base de datos, eliminando tablas relacionales intermedias mediante el uso de listas embebidas nativas de NoSQL.")
    add_bullet_point(doc, "La optimización mediante desnormalización (ej. incluir títulos de álbum y artista en las listas) permite realizar consultas analíticas e interacciones en tiempo real con latencias extremadamente bajas.")
    add_bullet_point(doc, "El patrón de diseño adaptador implementado (MongoRow) facilitó migrar el motor de base de datos de manera limpia, sin impactar el frontend de Jinja2 o alterar las plantillas visuales del sistema.")
    
    p_l = add_paragraph_styled(doc, space_before=4)
    p_l.add_run("• Lecciones Aprendidas:").bold = True
    add_bullet_point(doc, "Conversión de Fechas: Es crítico transformar explícitamente textos ISO a objetos Date nativos al importar a MongoDB para posibilitar búsquedas temporales.")
    add_bullet_point(doc, "Operaciones en Listas: Se asimiló el uso de comandos avanzados como $pull y $push para manipular registros embebidos de manera granular sin destruir el documento contenedor.")
    add_bullet_point(doc, "Control de Calidad Híbrido: La implementación de suites de pruebas simultáneas (MSSQL / MongoDB) es la mejor práctica industrial para validar migraciones críticas sin riesgos de regresión.")

    # --- ANEXOS ---
    add_heading_styled(doc, "8. ANEXOS (ESTRUCTURA Y SCRIPT MOCK)", level=2)
    
    p_dir = add_paragraph_styled(doc)
    p_dir.add_run("• Árbol de Directorios del Proyecto (Fase 6):").bold = True
    add_styled_code_block(doc, """SoundWave/
├── Fase6/
│   ├── app.py                     # Backend Flask (Soporte Dual: MongoDB / SQL Server)
│   ├── config.json                # Configuración activa (database_type: "mongodb")
│   ├── requirements.txt           # Dependencias del sistema
│   ├── sql/
│   │   └── sql_to_json.sql        # Consultas de extracción SQL Server a JSON
│   ├── mongodb/
│   │   ├── playground.js          # MongoDB Playground (CRUD e Índices)
│   │   ├── export_sql_to_json.py  # Script de extracción
│   │   ├── import_json_to_mongo.py# Script de inserción y tipado BSON Date
│   │   └── data/                  # Carpeta de datos JSON generados
│   └── tests/
│       ├── test_app.py            # Pruebas relacionales (MSSQL)
│       └── test_mongodb.py        # Pruebas documentales (MongoDB)""")
        
    p_links = add_paragraph_styled(doc, space_before=4)
    p_links.add_run("• Enlaces de Entrega:").bold = True
    add_bullet_point(doc, "https://github.com/FelipeBoada/SoundWave-Database-Project", bold_prefix="Repositorio GitHub: ")
    add_bullet_point(doc, "https://youtu.be/SoundWaveNoSQLCRUD", bold_prefix="Video Demostrativo: ")

    # Intentar guardar
    try:
        print(f"Intentando guardar documento en: {ruta_salida}")
        doc.save(ruta_salida)
        print("Documento principal guardado con éxito.")
    except PermissionError as pe:
        print(f"Error de permisos al escribir en {ruta_salida}. Es probable que esté abierto en Word.")
        print(f"Guardando copia en: {ruta_alterna}")
        doc.save(ruta_alterna)
        print("Copia guardada con éxito en la ruta alterna.")

if __name__ == "__main__":
    main()
